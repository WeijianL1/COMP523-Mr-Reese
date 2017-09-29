#!/usr/bin/python
# str(sys.argv[1])).text)
from docx import Document
import xlrd, datetime
from docx import Document
from collections import OrderedDict
import simplejson as json
import sys
import re

inFile=str(sys.argv[1])
wb=xlrd.open_workbook(inFile)
sh=wb.sheet_by_name("Sheet1")
document = Document()

# for rowx in range(sh.nrows):
for rowx in range(0,10):
	title=sh.row_values(rowx,0,1)[0]
	text=sh.row_values(rowx,1,2)[0]
	# text=text.replace("\n","")
	text=re.sub('\n+',' ',text)
	document.add_heading(title,1)
	document.add_paragraph(text)
document.save('demo.docx')
